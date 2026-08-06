"""
python-docx를 사용해 Word 문서를 생성하는 스크립트.
제목/본문 크기, 폰트(맑은 고딕), 메인 컬러를 설정값으로 조정할 수 있다.

주의: Word는 글자 크기를 px가 아닌 pt 단위로 다룬다. 요청한 40 / 20은
포인트(pt) 값으로 적용했다 (필요하면 TITLE_SIZE_PT, BODY_SIZE_PT 값을 조정).
"""

import argparse

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "맑은 고딕"
TITLE_SIZE_PT = 20
HEADING_SIZE_PT = 18
BODY_SIZE_PT = 12
DEFAULT_MAIN_COLOR = "2E74B5"  # 파란색 계열 (요청 색상으로 교체 가능)


def hex_to_rgbcolor(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        raise ValueError(f"올바른 hex 색상 코드가 아닙니다: {hex_color}")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def apply_east_asian_font(font_element, font_name: str) -> None:
    """python-docx의 font.name만으로는 한글 폰트가 적용되지 않아
    w:rFonts의 eastAsia 속성을 직접 지정해야 한다 (스타일/런 공통)."""
    rFonts = font_element.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        font_element.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str, size_pt: int, color: RGBColor = None, bold: bool = None) -> None:
    """문서 스타일(Title/Heading 1/Normal 등) 자체를 수정한다.
    스타일 단위로 지정하면 add_heading/add_paragraph로 만든 텍스트에
    자동으로 서식이 적용되고, Word의 개요/목차 기능도 그대로 유지된다."""
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    rPr = style.element.get_or_add_rPr()
    apply_east_asian_font(rPr, font_name)


def set_cell_background(cell, hex_color: str) -> None:
    """표 셀 배경색 지정 (python-docx 기본 API에 없어 XML로 직접 처리)."""
    hex_color = hex_color.lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_table_cell_text(cell, text: str, bold: bool = False, color: RGBColor = None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    apply_east_asian_font(run._element.get_or_add_rPr(), FONT_NAME)


def create_sample_document(
    output_path: str = "output.docx",
    title_text: str = "문서 제목",
    main_color: str = DEFAULT_MAIN_COLOR,
) -> None:
    doc = Document()
    color = hex_to_rgbcolor(main_color)

    # 스타일 단위로 폰트/크기/색상 설정 -> 이후 add_heading/add_paragraph에 자동 반영
    set_style_font(doc.styles["Normal"], FONT_NAME, BODY_SIZE_PT)
    set_style_font(doc.styles["Title"], FONT_NAME, TITLE_SIZE_PT, color=color, bold=True)
    set_style_font(doc.styles["Heading 1"], FONT_NAME, HEADING_SIZE_PT, color=color, bold=True)

    # 제목
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 소제목 + 본문
    doc.add_heading("1. 개요", level=1)
    doc.add_paragraph(
        "이 문서는 python-docx로 생성된 예시 문서입니다. "
        "필요한 내용으로 자유롭게 수정해서 사용하세요."
    )

    # 굵게/기울임 서식이 섞인 문단
    p = doc.add_paragraph()
    p.add_run("굵은 글씨").bold = True
    p.add_run(" 와 ")
    p.add_run("기울임 글씨").italic = True
    p.add_run(" 를 함께 사용할 수 있습니다.")

    # 글머리 기호 목록
    doc.add_heading("2. 목록 예시", level=1)
    for item in ["첫 번째 항목", "두 번째 항목", "세 번째 항목"]:
        doc.add_paragraph(item, style="List Bullet")

    # 표 (헤더 행은 메인 컬러 배경 + 흰 글씨)
    doc.add_heading("3. 표 예시", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["이름", "부서", "비고"]
    for cell, text in zip(table.rows[0].cells, headers):
        style_table_cell_text(cell, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_background(cell, main_color)

    rows_data = [
        ("홍길동", "개발팀", "-"),
        ("김철수", "기획팀", "-"),
    ]
    for name, dept, note in rows_data:
        row_cells = table.add_row().cells
        for cell, text in zip(row_cells, (name, dept, note)):
            style_table_cell_text(cell, text)

    # 페이지 여백 설정 예시
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    doc.save(output_path)
    print(f"문서가 저장되었습니다: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="python-docx 문서 생성 스크립트")
    parser.add_argument("--output", default="output.docx", help="저장할 파일 경로")
    parser.add_argument("--title", default="문서 제목", help="문서 제목 텍스트")
    parser.add_argument(
        "--color",
        default=DEFAULT_MAIN_COLOR,
        help="메인 컬러 hex 코드 (# 없이, 예: 2E74B5)",
    )
    args = parser.parse_args()

    create_sample_document(output_path=args.output, title_text=args.title, main_color=args.color)


if __name__ == "__main__":
    main()
