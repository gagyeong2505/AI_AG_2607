"""
스크랩한 기사 콘텐츠 JSON을 받아 서식이 적용된 Word(.docx) 보고서를 생성한다.
제목 40pt / 본문 20pt / 맑은 고딕 / 설정 가능한 메인 컬러를 스타일 단위로 적용하고,
기사별로 스크린샷 이미지를 본문 폭에 맞춰 삽입한다.

입력 JSON 형식은 ../references/content_schema.md 참고.

사용법:
    python generate_docx.py --input content.json --output report.docx [--color 2E74B5]
"""

import argparse
import json
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "맑은 고딕"
TITLE_SIZE_PT = 40
HEADING_SIZE_PT = 28
BODY_SIZE_PT = 20
DEFAULT_MAIN_COLOR = "2E74B5"
IMAGE_WIDTH_CM = 14


def hex_to_rgbcolor(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        raise ValueError(f"올바른 hex 색상 코드가 아닙니다: {hex_color}")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def apply_east_asian_font(font_element, font_name: str) -> None:
    """font.name만으로는 한글 폰트가 적용되지 않아 w:rFonts의 eastAsia 속성을 직접 지정한다."""
    rFonts = font_element.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        font_element.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str, size_pt: int, color: RGBColor = None, bold: bool = None) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    apply_east_asian_font(style.element.get_or_add_rPr(), font_name)


def add_italic_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    apply_east_asian_font(run._element.get_or_add_rPr(), FONT_NAME)


def add_article_section(doc: Document, article: dict) -> None:
    doc.add_heading(article["headline"], level=1)

    meta_bits = [article["source"]]
    if article.get("published"):
        meta_bits.append(article["published"])
    meta_bits.append(article["url"])
    add_italic_paragraph(doc, " · ".join(meta_bits))

    screenshot = article.get("screenshot")
    if screenshot:
        if os.path.exists(screenshot):
            doc.add_picture(screenshot, width=Cm(IMAGE_WIDTH_CM))
        else:
            print(f"경고: 스크린샷을 찾을 수 없어 건너뜁니다: {screenshot}")

    doc.add_paragraph(article["summary"])

    for point in article.get("key_points", []):
        doc.add_paragraph(point, style="List Bullet")


def build_document(content: dict, output_path: str, main_color: str) -> None:
    doc = Document()
    color = hex_to_rgbcolor(main_color)

    # 스타일 단위 적용 -> add_heading/add_paragraph로 만드는 모든 텍스트에 자동 반영되고
    # Word의 개요/목차 기능도 그대로 유지된다.
    set_style_font(doc.styles["Normal"], FONT_NAME, BODY_SIZE_PT)
    set_style_font(doc.styles["Title"], FONT_NAME, TITLE_SIZE_PT, color=color, bold=True)
    set_style_font(doc.styles["Heading 1"], FONT_NAME, HEADING_SIZE_PT, color=color, bold=True)

    title = doc.add_heading(content["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if content.get("subtitle"):
        sub = doc.add_paragraph(content["subtitle"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.italic = True

    for paragraph in content.get("intro", []):
        doc.add_paragraph(paragraph)

    for article in content.get("articles", []):
        add_article_section(doc, article)

    conclusion = content.get("conclusion", [])
    if conclusion:
        doc.add_heading("결론", level=1)
        for paragraph in conclusion:
            doc.add_paragraph(paragraph)

    sources = content.get("sources", [])
    if sources:
        doc.add_heading("참고 자료", level=1)
        for src in sources:
            doc.add_paragraph(src, style="List Bullet")

    section_fmt = doc.sections[0]
    section_fmt.left_margin = Cm(2)
    section_fmt.right_margin = Cm(2)

    doc.save(output_path)
    print(f"saved: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="스크랩 기사 콘텐츠 JSON으로 docx 보고서 생성")
    parser.add_argument("--input", required=True, help="콘텐츠 JSON 파일 경로")
    parser.add_argument("--output", required=True, help="저장할 .docx 파일 경로")
    parser.add_argument("--color", default=DEFAULT_MAIN_COLOR, help="메인 컬러 hex 코드 (# 없이, 예: 2E74B5)")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        content = json.load(f)

    output_dir = os.path.dirname(args.output)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    build_document(content, args.output, args.color)


if __name__ == "__main__":
    main()
