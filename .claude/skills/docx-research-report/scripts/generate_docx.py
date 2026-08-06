"""
JSON으로 정의된 콘텐츠를 받아 서식이 적용된 Word(.docx) 보고서를 생성한다.
제목 40pt / 본문 20pt / 맑은 고딕 / 설정 가능한 메인 컬러를 스타일 단위로 적용한다.

입력 JSON 형식은 ../references/content_schema.md 참고.

사용법:
    python generate_docx.py --input content.json --output report.docx [--color 2E74B5]
"""

import argparse
import json

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "맑은 고딕"
TITLE_SIZE_PT = 40
HEADING_SIZE_PT = 28
BODY_SIZE_PT = 20
DEFAULT_MAIN_COLOR = "2E74B5"


def hex_to_rgbcolor(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        raise ValueError(f"올바른 hex 색상 코드가 아닙니다: {hex_color}")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def apply_east_asian_font(font_element, font_name: str) -> None:
    """font.name만으로는 한글 폰트가 적용되지 않아 w:rFonts의 eastAsia 속성을 직접 지정한다."""
    rFonts = font_element.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        font_element.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str, size_pt: int, color: RGBColor = None, bold: bool = None) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    apply_east_asian_font(style.element.get_or_add_rPr(), font_name)


def set_cell_background(cell, hex_color: str) -> None:
    hex_color = hex_color.lstrip("#")
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_table_cell_text(cell, text: str, bold: bool = False, color: RGBColor = None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    apply_east_asian_font(run._element.get_or_add_rPr(), FONT_NAME)


def add_table(doc: Document, table_data: dict, main_color: str) -> None:
    columns = table_data["columns"]
    rows = table_data.get("rows", [])
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, columns):
        style_table_cell_text(cell, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_background(cell, main_color)
    for row in rows:
        row_cells = table.add_row().cells
        for cell, text in zip(row_cells, row):
            style_table_cell_text(cell, text)


def build_document(content: dict, output_path: str, main_color: str) -> None:
    doc = Document()
    color = hex_to_rgbcolor(main_color)

    # 스타일 단위 적용 -> add_heading/add_paragraph로 만드는 모든 텍스트에 자동 반영되고
    # Word의 개요/목차 기능도 그대로 유지된다.
    set_style_font(doc.styles["Normal"], FONT_NAME, BODY_SIZE_PT)
    set_style_font(doc.styles["Title"], FONT_NAME, TITLE_SIZE_PT, color=color, bold=True)
    set_style_font(doc.styles["Heading 1"], FONT_NAME, HEADING_SIZE_PT, color=color, bold=True)

    title = doc.add_heading(content["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if content.get("subtitle"):
        sub = doc.add_paragraph(content["subtitle"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.italic = True

    for section in content.get("sections", []):
        if section.get("heading"):
            doc.add_heading(section["heading"], level=1)

        for paragraph in section.get("paragraphs", []):
            doc.add_paragraph(paragraph)

        bullets = section.get("bullets", [])
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")

        table_data = section.get("table")
        if table_data:
            add_table(doc, table_data, main_color)

    sources = content.get("sources", [])
    if sources:
        doc.add_heading("참고 자료", level=1)
        for src in sources:
            doc.add_paragraph(src, style="List Bullet")

    section_fmt = doc.sections[0]
    section_fmt.left_margin = Cm(2)
    section_fmt.right_margin = Cm(2)

    doc.save(output_path)
    print(f"saved: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="JSON 콘텐츠로 docx 보고서 생성")
    parser.add_argument("--input", required=True, help="콘텐츠 JSON 파일 경로")
    parser.add_argument("--output", required=True, help="저장할 .docx 파일 경로")
    parser.add_argument("--color", default=DEFAULT_MAIN_COLOR, help="메인 컬러 hex 코드 (# 없이, 예: 2E74B5)")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        content = json.load(f)

    build_document(content, args.output, args.color)


if __name__ == "__main__":
    main()
